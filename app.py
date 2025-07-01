import pandas as pd
import streamlit as st
import seaborn as sns
import matplotlib.pyplot as plt
import numpy as np
from sklearn.linear_model import LinearRegression
from io import BytesIO
from docx import Document
import os

# =============================
# ⚙️ Configuration page + thème + police
# =============================
st.set_page_config(
    page_title="🫒 Suivi de la fermentation des olives",
    page_icon="🫒",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Injecter style CSS (thème clair pistache + police Poppins)
st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Poppins&display=swap" rel="stylesheet">
<style>
html, body, [class*="css"]  {
    font-family: 'Poppins', sans-serif !important;
    background-color: #f4fff0 !important;  /* pistache clair */
    color: #1b3a1a !important; /* vert foncé */
}
.stApp {
    background-color: #f4fff0 !important;
}
h1, h2, h3, h4, h5 {
    color: #2e7d32 !important; /* vert un peu plus foncé */
}
section.main > div {
    background-color: white !important;
    border-radius: 15px;
    padding: 20px;
    box-shadow: 0 2px 8px rgba(46, 125, 50, 0.2);
}
</style>
""", unsafe_allow_html=True)

# =============================
# 🖼️ Bannière d'en-tête avec logos + titre
# =============================

col1, col2, col3 = st.columns([1,6,1])

with col1:
    logo_entreprise_path = "C:/Users/admin/Documents/projet fermentation/logo_entreprise.jpeg"
    if os.path.exists(logo_entreprise_path):
        st.image(logo_entreprise_path, width=100)
    else:
        st.warning("Logo entreprise introuvable")

with col2:
    st.markdown("<h1 style='text-align:center; margin-bottom:0;'>🫒 Suivi de la fermentation des olives</h1>", unsafe_allow_html=True)
    st.markdown("""
    <p style='text-align:center; font-size:16px; margin-top:5px;'>
    Plateforme professionnelle pour le suivi, l'analyse et la prédiction des données physico-chimiques<br>
    des cuves d'olives en fermentation.
    </p>
    """, unsafe_allow_html=True)

with col3:
    logo_iav_path = "C:/Users/admin/Documents/projet fermentation/logo_iav.png"
    if os.path.exists(logo_iav_path):
        st.image(logo_iav_path, width=100)
    else:
        st.warning("Logo IAV introuvable")

st.markdown("---")


# =============================
# 📂 Chargement fichier
# =============================
uploaded_file = st.sidebar.file_uploader("📁 Charger un fichier Excel ou CSV", type=["xlsx", "csv"])

if uploaded_file is not None:
    try:
        if uploaded_file.name.endswith(".csv"):
            try:
                df = pd.read_csv(uploaded_file, encoding="utf-8")
            except UnicodeDecodeError:
                df = pd.read_csv(uploaded_file, encoding="latin1")
        else:
            df = pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"Erreur de chargement du fichier : {e}")
        st.stop()
else:
    st.warning("📌 Veuillez charger un fichier Excel (.xlsx) ou CSV dans la barre latérale.")
    st.stop()

# =============================
# ✅ Préparation / nettoyage
# =============================
df.rename(columns=lambda x: x.strip(), inplace=True)

colonnes_necessaires = ["N°Cuve", "Date de remplissage", "PH", "AL"]
for col in colonnes_necessaires:
    if col not in df.columns:
        st.error(f"❌ Colonne manquante : {col}")
        st.stop()

df["Date de remplissage"] = pd.to_datetime(df["Date de remplissage"], errors="coerce")

if "La date" in df.columns:
    df["La date"] = pd.to_datetime(df["La date"], errors="coerce")
    df["Jours_apres_remplissage"] = (df["La date"] - df["Date de remplissage"]).dt.days
else:
    df["Jours_apres_remplissage"] = (pd.Timestamp.today() - df["Date de remplissage"]).dt.days

df = df[df["Jours_apres_remplissage"].notna() & (df["Jours_apres_remplissage"] >= 0)]
df["N°Cuve"] = df["N°Cuve"].astype(str)

# =============================
# 🔢 Sidebar organisée
# =============================
st.sidebar.title("⚙️ Paramètres de filtration")
cuves_valides = df["N°Cuve"].unique().tolist()
cuves_selectionnees = st.sidebar.multiselect("💡 Cuves à afficher", sorted(cuves_valides), default=cuves_valides)
min_jour = int(df["Jours_apres_remplissage"].min())
max_jour = int(df["Jours_apres_remplissage"].max())
jours_selectionnes = st.sidebar.slider("🗓️ Intervalle de jours", min_value=min_jour, max_value=max_jour, value=(min_jour, max_jour))
st.sidebar.markdown("---")
menu = st.sidebar.radio("📌 Navigation", ["Accueil", "Graphiques", "Rapport", "Prédiction", "Résumé Qualité"])

# =============================
# 🔢 Filtrage données
# =============================
df_filtre = df[df["N°Cuve"].isin(cuves_selectionnees)]
df_filtre = df_filtre[df_filtre["Jours_apres_remplissage"].between(*jours_selectionnes)].reset_index(drop=True)

# =============================
# 🧮 Mini Dashboard KPI
# =============================
total_analyses = df_filtre.shape[0]
nb_alertes = df_filtre[(df_filtre["PH"] > 4.5) | (df_filtre["AL"] <= 0.8)].shape[0]
nb_conformes = total_analyses - nb_alertes
jours_analysee = max_jour - min_jour + 1

st.markdown(f"""
<div style="display: flex; justify-content: space-around; background-color: white; border-radius: 15px; padding: 20px; box-shadow: 0 2px 12px rgba(46, 125, 50, 0.25); margin-bottom: 25px;">
    <div style="text-align: center;">
        <h3 style="color:#2e7d32;">📦 {len(cuves_selectionnees)}</h3>
        <p>Cuves sélectionnées</p>
    </div>
    <div style="text-align: center;">
        <h3 style="color:#2e7d32;">📅 {jours_analysee}</h3>
        <p>Jours analysés</p>
    </div>
    <div style="text-align: center;">
        <h3 style="color:#c62828;">⚠️ {nb_alertes}</h3>
        <p>Alertes détectées</p>
    </div>
</div>
""", unsafe_allow_html=True)

# =============================
# 📌 Menu Navigation
# =============================

if menu == "Accueil":
    st.header("📋 Aperçu des données filtrées")
    st.markdown("""
    ### ℹ️ Guide utilisateur
    - 📁 Chargez votre fichier Excel ou CSV contenant les données des cuves.
    - 🧪 Sélectionnez les cuves et la période à analyser dans la barre latérale.
    - 📈 Visualisez l'évolution des paramètres physico-chimiques.
    - 📄 Générez automatiquement un rapport complet.
    """)
    st.dataframe(df_filtre.head(20))

elif menu == "Graphiques":
    st.header("📈 Visualisation graphique")
    params = st.multiselect("Choisissez un ou plusieurs paramètres à tracer",
                            ["°Be", "% Na Cl", "PH", "AL", "AC", "T°C"],
                            default=["PH"])
    for param in params:
        if param not in df_filtre.columns:
            st.warning(f"Paramètre {param} non trouvé dans les données.")
            continue
        st.subheader(f"Paramètre : {param}")
        fig, ax = plt.subplots(figsize=(10, 6))
        sns.lineplot(data=df_filtre, x="Jours_apres_remplissage", y=param,
                     hue="N°Cuve", marker="o", ax=ax)
        ax.set_title(f"Évolution de {param}")
        ax.set_xlabel("Jours")
        ax.set_ylabel(param)
        ax.grid(True)
        st.pyplot(fig)

elif menu == "Rapport":
    st.header("📄 Rapport automatique")
    rapport = []
    for _, row in df_filtre.iterrows():
        cuve = row["N°Cuve"]
        date = row.get("La date", None) or row.get("Date vidange ")
        date_str = date.strftime("%d/%m/%Y") if pd.notna(date) else "Date manquante"
        ph = row["PH"]
        al = row["AL"]

        bloc = [f"**🧪 Cuve {cuve} | Date : {date_str}**"]
        if pd.isna(ph):
            bloc.append("⚠️ pH manquant")
        elif ph <= 4.2:
            bloc.append("✔️ pH conforme (≤ 4.2)")
        elif ph < 4.5:
            bloc.append("⚠️ pH > 4.2")
        else:
            bloc.append("❌ pH ≥ 4.5")

        if not pd.isna(al):
            bloc.append("✔️ AL analysée" if al > 0.8 else "❌ AL ≤ 0.8")
        rapport.append("\n".join(bloc))

    st.markdown("\n\n".join(rapport))

    if st.button("📄 Exporter le rapport Word"):
        doc = Document()
        doc.add_heading("Rapport de Suivi de Fermentation", 0)
        for bloc in rapport:
            doc.add_paragraph(bloc, style="List Bullet")
        buffer = BytesIO()
        doc.save(buffer)
        st.download_button("📂 Télécharger le rapport Word", buffer.getvalue(), file_name="rapport_fermentation.docx")

elif menu == "Prédiction":
    st.header("🔮 Prédiction de pH et AL")
    for param in ["PH", "AL"]:
        st.subheader(f"🔬 Prédiction : {param}")
        if param not in df_filtre.columns:
            st.warning(f"Paramètre {param} non trouvé.")
            continue
        df_pred = df_filtre[["Jours_apres_remplissage", param]].dropna()
        if df_pred.shape[0] < 5:
            st.warning(f"Pas assez de données pour {param}.")
            continue

        X = df_pred["Jours_apres_remplissage"].values.reshape(-1, 1)
        y = df_pred[param].values
        model = LinearRegression().fit(X, y)
        jours_future = np.arange(0, 26).reshape(-1, 1)
        predictions = model.predict(jours_future)
        std_err = np.std(y - model.predict(X))

        fig, ax = plt.subplots(figsize=(8, 5))
        ax.scatter(X.flatten(), y, label="Données", color="blue")
        ax.plot(jours_future.flatten(), predictions, label="Prédiction", color="red")
        ax.fill_between(jours_future.flatten(), predictions - std_err, predictions + std_err,
                        color="red", alpha=0.2, label="Zone d'incertitude")
        ax.set_title(f"Tendance prédite : {param}")
        ax.set_xlabel("Jours")
        ax.set_ylabel(param)
        ax.grid(True)
        ax.legend()
        st.pyplot(fig)
        st.markdown(f"📈 **Modèle :** {param} = {model.coef_[0]:.3f} * jours + {model.intercept_:.3f}")

elif menu == "Résumé Qualité":
    st.header("📊 Résumé de la Qualité")
    st.markdown(f"""
    - 🔍 Nombre total d’analyses filtrées : **{total_analyses}**
    - ✅ Conformes : **{nb_conformes}**
    - ⚠️ Alertes : **{nb_alertes}**
    """)

    if nb_alertes > 0:
        st.error("❗ Des valeurs non conformes ont été détectées.")
    else:
        st.success("✅ Toutes les valeurs sont dans les normes.")

    cuves_alertes = df_filtre[df_filtre["PH"] > 4.5]["N°Cuve"].unique().tolist()
    if cuves_alertes:
        st.warning(f"➡️ Cuves à surveiller (pH > 4.5) : {', '.join(cuves_alertes)}")
    else:
        st.info("Aucune cuve avec un pH > 4.5.")

# =============================
# Pied de page discret
# =============================
st.markdown("""
<hr style='margin-top: 50px;'>
<p style='text-align: center; font-size: 0.8em; color: grey'>
Développé par Fatima Ezzahra Lafdali – Projet de suivi en fermentation | IAV 2025
</p>
""", unsafe_allow_html=True)
